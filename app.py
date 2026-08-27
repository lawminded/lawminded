import json
import os
import re
import sqlite3
import base64
import hmac
import threading
import time
from io import BytesIO
from datetime import date, datetime, timedelta, timezone, timedelta
from functools import wraps, lru_cache
from flask import (Flask, render_template, request, redirect, url_for,
                   has_request_context,
                   session, flash, jsonify, send_from_directory, send_file, abort)
from flask_mail import Mail, Message
from itsdangerous import URLSafeSerializer, URLSafeTimedSerializer, BadSignature, SignatureExpired
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from werkzeug.middleware.proxy_fix import ProxyFix
from flask_wtf import CSRFProtect
from flask_wtf.csrf import CSRFError
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_talisman import Talisman
import bleach
from dotenv import load_dotenv

# Load .env BEFORE importing database, which reads DATABASE_PATH at import time.
# (If load_dotenv runs after that import, the custom DB path is silently ignored.)
load_dotenv()

from database import (get_db, init_db, seed_articles, seed_documents, seed_formats,
                      apply_content_migrations, DB_PATH)
import content as C
import formats as F
from seo_meta import (SEO_DESCRIPTIONS, SEO_TITLES, SEARCH_META_CHANGED,
                      INTERNAL_LINKS, RETIRED_ARTICLES)

# Production flag: enables HTTPS enforcement, HSTS, and Secure cookies on the live
# server. Stays off locally so http://localhost development still works.
IS_PROD = os.getenv('PRODUCTION', 'false').lower() in ('1', 'true', 'yes')

app = Flask(__name__)

# The session cookie is what marks an admin as logged in, so the signing key is
# effectively the admin credential. In production it MUST come from .env: if the
# file ever fails to load, falling back to a known string published in this repo
# would let anyone mint their own "admin_logged_in" cookie. Fail closed instead.
_secret = os.getenv('SECRET_KEY')
if IS_PROD and not _secret:
    raise RuntimeError(
        'SECRET_KEY is not set. Refusing to start in production with a default '
        'signing key — set SECRET_KEY in .env.'
    )
app.secret_key = _secret or 'dev-secret-change-in-production'

# Behind a reverse proxy (nginx on the Oracle VM), trust the X-Forwarded-* headers
# so request.scheme is 'https'. Without this, Talisman's force_https would
# redirect-loop forever when the app sits behind nginx.
#
# Only x_for/x_proto are trusted: nginx sets both itself, so a client cannot forge
# them. It does NOT set X-Forwarded-Host or -Port, so those arrive straight from
# the client — trusting them would let anyone rewrite request.host (poisoned
# absolute URLs, redirects and cache entries). nginx now pins both headers too,
# so this is belt and braces.
if IS_PROD:
    app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=0, x_port=0)

# ── Hardened session / request config ──
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['SESSION_COOKIE_SECURE'] = IS_PROD
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16 MB — room for admin Word (.docx) uploads
app.config['WTF_CSRF_TIME_LIMIT'] = None            # token valid for the session
# Admin sessions expire after 8 idle hours (Flask refreshes the cookie on every
# request), so a forgotten login on a shared machine does not stay valid forever.
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=8)

# Mail configuration — provider-agnostic via env (defaults to Gmail SMTP).
# Set MAIL_SERVER/PORT for a custom-domain mailbox (e.g. Zoho, Titan, Workspace).
app.config['MAIL_SERVER'] = os.getenv('MAIL_SERVER', 'smtp.gmail.com')
app.config['MAIL_PORT'] = int(os.getenv('MAIL_PORT', '587'))
app.config['MAIL_USE_TLS'] = os.getenv('MAIL_USE_TLS', 'true').lower() in ('1', 'true', 'yes')
app.config['MAIL_USE_SSL'] = os.getenv('MAIL_USE_SSL', 'false').lower() in ('1', 'true', 'yes')
app.config['MAIL_USERNAME'] = os.getenv('MAIL_USERNAME')
app.config['MAIL_PASSWORD'] = os.getenv('MAIL_PASSWORD')
# Sender can differ from the login user (defaults to the login user).
app.config['MAIL_DEFAULT_SENDER'] = os.getenv('MAIL_DEFAULT_SENDER', os.getenv('MAIL_USERNAME'))

mail = Mail(app)

# ── CSRF protection on every POST form ──
csrf = CSRFProtect(app)

# ── Rate limiting (brute-force + spam protection) ──
limiter = Limiter(
    get_remote_address,
    app=app,
    default_limits=['300 per hour'],
    storage_uri='memory://',
)

# ── Security headers + Content-Security-Policy ──
# script-src and connect-src used to be a blanket `https:`, which let injected
# markup pull a payload from any HTTPS host anywhere. Every script this site
# actually loads is either same-origin or Google AdSense, so the policy now names
# those origins and nothing else.
#
# The AdSense origins are only added while ADS_ENABLED is on (see below), so
# while ads are switched off the policy is same-origin only. `unsafe-inline`
# stays because templates carry inline <script> and JSON-LD blocks; all
# admin-entered HTML is bleach-sanitised before it is ever rendered.
_GOOGLE_ADS_SCRIPT = [
    'https://pagead2.googlesyndication.com',
    'https://tpc.googlesyndication.com',
    'https://partner.googleadservices.com',
    'https://adservice.google.com',
    'https://www.googletagservices.com',
]
_GOOGLE_ADS_CONNECT = [
    'https://pagead2.googlesyndication.com',
    'https://googleads.g.doubleclick.net',
    'https://csi.gstatic.com',
]


def _build_csp(ads_on):
    return {
        'default-src': "'self'",
        'script-src': ["'self'", "'unsafe-inline'"] + (_GOOGLE_ADS_SCRIPT if ads_on else []),
        'style-src': ["'self'", "'unsafe-inline'", 'https://fonts.googleapis.com'],
        'font-src': ["'self'", 'https://fonts.gstatic.com', 'data:'],
        'img-src': ["'self'", 'data:', 'https:'],   # ad creatives + remote article images
        'frame-src': (['https://*.googlesyndication.com', 'https://*.google.com',
                       'https://*.doubleclick.net'] if ads_on else ["'none'"]),
        'connect-src': ["'self'"] + (_GOOGLE_ADS_CONNECT if ads_on else []),
        'object-src': "'none'",
        'base-uri': "'self'",
        'frame-ancestors': "'self'",
        'form-action': "'self'",
    }


# Site-wide ad kill-switch. Declared here (rather than with the other AdSense
# config further down) because the CSP above widens only when ads are on.
# While False, no ad script loads and every ad zone renders nothing.
ADS_ENABLED = False

Talisman(
    app,
    content_security_policy=_build_csp(ADS_ENABLED),
    force_https=IS_PROD,
    strict_transport_security=IS_PROD,
    session_cookie_secure=IS_PROD,
    referrer_policy='strict-origin-when-cross-origin',
)


@app.errorhandler(CSRFError)
def handle_csrf_error(e):
    return render_template('error.html', code=400,
                           message='Your session expired. Please go back and try again.'), 400


# ── HTML sanitisation for admin-entered article content ──
ALLOWED_TAGS = ['h2', 'h3', 'h4', 'p', 'ul', 'ol', 'li', 'strong', 'em', 'b', 'i',
                'u', 'a', 'br', 'blockquote', 'code', 'pre', 'span', 'hr', 'table',
                'thead', 'tbody', 'tr', 'th', 'td']
ALLOWED_ATTRS = {'a': ['href', 'title', 'rel', 'target'], 'span': ['class'],
                 'td': ['colspan', 'rowspan'], 'th': ['colspan', 'rowspan']}


def sanitize_html(raw):
    return bleach.clean(raw or '', tags=ALLOWED_TAGS, attributes=ALLOWED_ATTRS, strip=True)

CONTACT_RECEIVER = os.getenv('CONTACT_RECEIVER', os.getenv('MAIL_USERNAME'))
# Admin credentials live in .env (ADMIN_USERNAME + ADMIN_PW_HASH_B64), never in
# the database — so a database reset can NEVER revert access to a default password.
# There is intentionally NO hardcoded default: if credentials are unset, the admin
# login fails closed (no access until credentials are configured on the server).
ENV_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.env')

# Google AdSense — set ADSENSE_CLIENT in .env to your publisher id (ca-pub-XXXX).
# Ad slot ids are configured per-placement in ADSENSE_SLOTS below.
# The ADS_ENABLED kill-switch is declared further up, next to the CSP it widens.
# Turned OFF while re-applying for AdSense approval — flip it back to True there
# to restore ads (the config below is preserved either way).
ADSENSE_CLIENT = os.getenv('ADSENSE_CLIENT', '')
ADSENSE_SLOTS = {
    'top': os.getenv('ADSENSE_SLOT_TOP', ''),
    'mid': os.getenv('ADSENSE_SLOT_MID', ''),
    'bottom': os.getenv('ADSENSE_SLOT_BOTTOM', ''),
    'article_top': os.getenv('ADSENSE_SLOT_ARTICLE_TOP', ''),
    'article_bottom': os.getenv('ADSENSE_SLOT_ARTICLE_BOTTOM', ''),
}

# Google Search Console — paste the "content" value from the HTML-tag
# verification method into GOOGLE_SITE_VERIFICATION in .env. When set, base.html
# renders <meta name="google-site-verification" ...> so Google can verify the site.
GOOGLE_SITE_VERIFICATION = os.getenv('GOOGLE_SITE_VERIFICATION', '')

VALID_EMAIL_RE = re.compile(r'^[^\s@]+@[^\s@]+\.[^\s@]+$')

# Public base URL of the live site (used for canonical tags, sitemap, OG, JSON-LD).
SITE_URL = os.getenv('SITE_URL', 'https://lawminded.in').rstrip('/')

# ── Named author (E-E-A-T) ──────────────────────────────────────────────────
# Legal content is YMYL, where Google weights demonstrable, *named* authorship
# most heavily. This block drives the article byline and the Person schema.
#
# Deliberately minimal by the owner's instruction: name and the held
# qualification, nothing else. There is no author bio page — it was removed in
# Aug 2026 and /author/<slug> now 301s to /about. Do NOT add a bio page, photo,
# jobTitle, or sameAs links to personal profiles/sites, and do not make the
# byline link anywhere. Their absence is a decision, not an oversight, and SEO
# advice to "strengthen the author entity" does not override it.
AUTHOR = {
    'name': 'Piyush Kundnani',
    'slug': 'piyush-kundnani',
    'credential': 'B.Com',   # genuinely held — byline + schema hasCredential
}

# The nine landmark-judgment briefs are hand-written editorial content that
# shipped together in the initial commit. Articles carry per-row timestamps from
# the database; these are static, so one honest publication date drives their
# Article schema and sitemap lastmod. Bump it if the briefs are ever rewritten.
JUDGMENTS_PUBLISHED = '2026-06-21'

CATEGORY_MAP = C.CATEGORY_MAP

# Signed, tamper-proof tokens for one-click newsletter unsubscribe links.
# itsdangerous ships with Flask, so no new dependency is needed.
_unsub_serializer = URLSafeSerializer(app.secret_key, salt='lm-unsubscribe')


def unsubscribe_url(email):
    """Absolute, signed unsubscribe link for inclusion in outgoing emails.

    url_for needs a request context, and the biggest caller of this — the
    announcement sent when an article is published — runs on a background
    thread that has only an app context. Without the fallback below, every
    subscriber email fails on building the link rather than on the mail
    server: a confusing way to discover you have sent nothing.
    """
    token = _unsub_serializer.dumps(email)
    if has_request_context():
        path = url_for("unsubscribe")
    else:
        with app.test_request_context():
            path = url_for("unsubscribe")
    return f'{SITE_URL}{path}?token={token}'


# The site's readers, its deadlines and its article dates are all Indian. The
# server runs UTC and so does SQLite's CURRENT_TIMESTAMP, which puts anything
# published after 18:30 UTC — i.e. after midnight in Delhi — on the previous day.
IST = timezone(timedelta(hours=5, minutes=30))


def ist_now():
    """Current IST wall-clock in the format the articles table stores."""
    return datetime.now(IST).strftime('%Y-%m-%d %H:%M:%S')


def ist_today():
    return datetime.now(IST).strftime('%Y-%m-%d')


# Same idea for the weekly draft-approval links sent to Telegram. The scheduled
# writer inserts an article as published=0 (already invisible to /article) and
# messages a signed preview link; approving it is a POST from that page.
_draft_serializer = URLSafeSerializer(app.secret_key, salt='lm-draft')


def draft_url(slug):
    """Absolute, signed preview link for an unpublished draft."""
    return f'{SITE_URL}/draft/{slug}?t={_draft_serializer.dumps(slug)}'


# ─── Public-form abuse defence ───────────────────────────────────────────────
# A form-spam campaign ("RobertMom" — short "what is your price" messages in
# rotating languages) put 3,300 entries through /contact over nine days in Aug
# 2026 from two IPs in 80.94.95.0/24. It beat the honeypot and CSRF by loading
# the real page first and filling only the visible fields, and never tripped
# the old "6 per minute" cap because it paced itself at about one a minute.
#
# The damage was not the noise: every message also sent an acknowledgement
# email to the attacker-supplied address, so the site emitted ~3,300 unsolicited
# emails to ~390 harvested third-party addresses. That is backscatter, and it
# puts the sending domain's reputation (and the newsletter) at risk.
#
# Three independent layers below, because any one of them can be worked around:
#   1. a dwell-time token — the form must have been open for a few seconds,
#   2. a hard per-IP rate limit on the route itself,
#   3. content heuristics for the payloads bots actually send.
_form_serializer = URLSafeTimedSerializer(app.secret_key, salt='lm-form')

MIN_FORM_SECONDS = 3        # humans never submit faster than this
MAX_FORM_SECONDS = 60 * 60 * 6   # a stale tab shouldn't hard-fail either


def form_token():
    """Signed 'this form was rendered now' token, exposed to templates."""
    return _form_serializer.dumps('f')


def form_token_ok(token):
    """True when the token is ours and the form was open long enough."""
    if not token:
        return False
    try:
        _form_serializer.loads(token, max_age=MAX_FORM_SECONDS)
    except (BadSignature, SignatureExpired):
        return False
    try:
        age = time.time() - _form_serializer.loads(token, max_age=MAX_FORM_SECONDS,
                                                  return_timestamp=True)[1].timestamp()
    except Exception:
        return False
    return age >= MIN_FORM_SECONDS


_SPAM_MARKERS = (
    'know your price', 'knew your price', 'wanted to know your price',
    'vestri pretium', 'tuo prezzo', 'kumukūʻai', 'din pris', 'su precio',
    'ihren preis', 'votre prix', 'вашу цену',
)
_URL_RE = re.compile(r'(https?://|www\.|\[url|<a\s+href)', re.I)


def looks_like_spam(name, email, query):
    """Cheap content heuristics for the payloads bots actually send.

    Deliberately narrow — a false positive here silently drops a real person's
    message, which is far worse than letting one spam entry through.
    """
    blob = f'{name} {query}'.lower()
    if _URL_RE.search(query or ''):
        return True                      # contact form has no reason to carry links
    if any(m in blob for m in _SPAM_MARKERS):
        return True
    # Deliberately NOT matching run-together capitalised names ("RobertMom").
    # It looks like a clean signature until you remember McDonald, MacLeod,
    # DeSouza and D'Souza are real surnames — that rule silently binned real
    # enquiries. The dwell-time token and the per-IP cap already cover the
    # generic case; this function only needs to catch obvious payloads.
    return False


app.jinja_env.globals['form_token'] = form_token


# ─── Branded transactional email ─────────────────────────────────────────────
# A single shell keeps every outgoing email on-brand (logo, palette) and
# improves deliverability: real plain-text part, List-Unsubscribe header, and
# an inline (cid:) logo so it shows even when remote images are blocked.
EMAIL_LOGO = 'static/img/logo-horizontal.png'


def _email_html(heading, body_html, unsub=None):
    foot_unsub = (
        '<p style="margin:14px 0 0;font:400 12px/1.6 Arial,sans-serif;color:#A99B78;">'
        'You are receiving this because you subscribed at Law Minded. '
        f'<a href="{unsub}" style="color:#8A5E07;">Unsubscribe</a>.</p>'
    ) if unsub else ''
    return (
        '<!doctype html><html lang="en"><body style="margin:0;padding:0;background:#F4F0E6;">'
        '<table role="presentation" width="100%" cellpadding="0" cellspacing="0" '
        'style="background:#F4F0E6;padding:28px 12px;font-family:Arial,Helvetica,sans-serif;">'
        '<tr><td align="center">'
        '<table role="presentation" width="100%" cellpadding="0" cellspacing="0" '
        'style="max-width:540px;background:#ffffff;border:1px solid #ECE6D8;border-radius:14px;overflow:hidden;">'
        '<tr><td style="height:4px;background:#E8A020;font-size:0;line-height:0;">&nbsp;</td></tr>'
        '<tr><td style="padding:26px 34px 0;">'
        '<img src="cid:logo" alt="LAW MiNDED" width="168" style="display:block;height:auto;border:0;outline:none;text-decoration:none;">'
        '<p style="margin:10px 0 0;font:600 10px/1 Arial,sans-serif;letter-spacing:.16em;'
        'text-transform:uppercase;color:#B7AC8E;">Where complexity meets clarity</p>'
        '</td></tr>'
        '<tr><td style="padding:20px 34px 28px;color:#33302A;">'
        f'<h1 style="margin:0 0 14px;font:700 20px/1.3 Georgia,\'Times New Roman\',serif;color:#1C1B16;">{heading}</h1>'
        f'{body_html}'
        '<hr style="border:0;border-top:1px solid #ECE6D8;margin:22px 0 12px;">'
        '<p style="margin:0;font:400 12px/1.6 Arial,sans-serif;color:#A99B78;">'
        'Law&nbsp;Minded — India\'s free legal-awareness platform.</p>'
        f'{foot_unsub}'
        '</td></tr></table></td></tr></table></body></html>'
    )


def send_branded_email(subject, recipients, heading, body_html, body_text, unsub=None):
    """Build + send a branded HTML email (with plain-text alt + inline logo)."""
    msg = Message(subject=subject, recipients=recipients)
    msg.body = body_text
    msg.html = _email_html(heading, body_html, unsub)
    if CONTACT_RECEIVER:
        msg.reply_to = CONTACT_RECEIVER
    if unsub:
        # RFC 2369 + RFC 8058 one-click unsubscribe — a strong inbox-placement signal.
        msg.extra_headers = {
            'List-Unsubscribe': f'<{unsub}>',
            'List-Unsubscribe-Post': 'List-Unsubscribe=One-Click',
        }
    with app.open_resource(EMAIL_LOGO) as f:
        msg.attach('logo.png', 'image/png', f.read(), 'inline',
                   headers={'Content-ID': '<logo>', 'X-Attachment-Id': 'logo'})
    mail.send(msg)


def _log_email(db, recipient, subject, kind, slug, status, error=None):
    db.execute(
        'INSERT INTO email_log (recipient, subject, kind, article_slug, status, error) '
        'VALUES (?,?,?,?,?,?)',
        (recipient, subject, kind, slug, status, (error or '')[:500]))


def mail_subscribers(subject, heading, body_html, body_text, kind, slug=None):
    """Send one message to every subscriber, one SMTP conversation each, and
    record the outcome per recipient.

    Individually rather than one message with everyone in the To line, for two
    reasons: subscribers must not see each other's addresses, and each person's
    unsubscribe link has to be their own. The per-recipient log exists because
    SMTP tells you nothing after the fact — 'it was sent' is not a record of who
    received what.

    Returns (sent, failed). Never raises: a mail server having a bad afternoon
    must not take down whatever called this.
    """
    db = get_db()
    people = db.execute('SELECT email, name FROM subscribers').fetchall()
    sent = failed = 0
    for person in people:
        addr = (person['email'] or '').strip()
        if not addr:
            continue
        try:
            send_branded_email(subject, [addr], heading, body_html, body_text,
                               unsub=unsubscribe_url(addr))
            _log_email(db, addr, subject, kind, slug, 'sent')
            sent += 1
        except Exception as e:                       # noqa: BLE001 — see docstring
            _log_email(db, addr, subject, kind, slug, 'failed', repr(e))
            failed += 1
            app.logger.warning('newsletter to %s failed: %r', addr, e)
    db.commit()
    db.close()
    return sent, failed


def announce_article(slug):
    """Tell subscribers about a newly published article. Called after the owner
    publishes, in a background thread so the tap returns immediately."""
    db = get_db()
    row = db.execute(
        'SELECT title, slug, summary, read_time, category FROM articles '
        'WHERE slug=? AND published=1', (slug,)).fetchone()
    db.close()
    if not row:
        return 0, 0

    url = f'{SITE_URL}/article/{row["slug"]}'
    cat = CATEGORY_MAP.get(row['category'], 'Guide')
    summary = (row['summary'] or '').strip()

    body_html = (
        f'<p style="margin:0 0 14px;">We have just published something new on '
        f'Law Minded, and it looked worth putting in front of you.</p>'
        f'<p style="margin:0 0 6px;font-size:13px;letter-spacing:.06em;'
        f'text-transform:uppercase;color:#8a6412;">{cat}</p>'
        f'<p style="margin:0 0 10px;font-size:19px;font-weight:600;line-height:1.35;">'
        f'{row["title"]}</p>'
        f'<p style="margin:0 0 18px;">{summary}</p>'
        f'<p style="margin:0 0 20px;"><a href="{url}" '
        f'style="background:#E8A020;color:#1a1a1a;padding:11px 22px;border-radius:4px;'
        f'text-decoration:none;font-weight:600;display:inline-block;">Read it '
        f'&rarr;</a></p>'
        f'<p style="margin:0;color:#6b6b6b;font-size:14px;">'
        f'{row["read_time"] or "A few minutes"} &middot; free to read, no sign-up.</p>'
    )
    body_text = (
        f'We have just published something new on Law Minded.\n\n'
        f'{row["title"]}\n\n{summary}\n\nRead it: {url}\n'
    )
    return mail_subscribers(f'New on Law Minded: {row["title"]}',
                            'A new guide is up', body_html, body_text,
                            kind='new-article', slug=row['slug'])


def announce_article_async(slug):
    """Publishing should feel instant. The send happens behind the redirect."""
    def _run():
        with app.app_context():
            try:
                s, f = announce_article(slug)
                app.logger.info('announced %s to subscribers: %d sent, %d failed',
                                slug, s, f)
            except Exception as e:                   # noqa: BLE001
                app.logger.exception('announcing %s failed: %r', slug, e)
    threading.Thread(target=_run, daemon=True).start()


# ─── Context Processor (globals available in every template) ─────────────────

def asset_version(rel_path):
    """Cache-busting token (file mtime) so browsers always fetch the latest CSS/JS."""
    try:
        return int(os.path.getmtime(os.path.join(app.static_folder, rel_path)))
    except OSError:
        return 1


@app.template_filter('humandate')
def humandate(value):
    """Render a stored timestamp as day-first, e.g. '21 Jun 2026'."""
    if not value:
        return ''
    s = str(value)[:10]
    try:
        from datetime import datetime as _dt
        return _dt.strptime(s, '%Y-%m-%d').strftime('%d %b %Y')
    except ValueError:
        return s


@app.template_filter('iso8601')
def iso8601(value, offset='+05:30'):
    """Render a stored timestamp as valid ISO-8601 for structured data,
    e.g. '2026-06-27T15:26:32+05:30'. Google requires the 'T' separator and a
    timezone designator; the raw SQLite 'YYYY-MM-DD HH:MM:SS' form is invalid.
    The wall-clock value is kept as-is and tagged with the site offset (IST),
    consistent with how humandate renders the same value."""
    if not value:
        return ''
    from datetime import datetime as _dt
    s = str(value).strip().replace('T', ' ')[:19]
    for fmt in ('%Y-%m-%d %H:%M:%S', '%Y-%m-%d %H:%M', '%Y-%m-%d'):
        try:
            return _dt.strptime(s, fmt).strftime('%Y-%m-%dT%H:%M:%S') + offset
        except ValueError:
            continue
    return str(value)


# Google renders roughly 60 characters of a <title> and cuts the rest mid-word.
TITLE_MAX = 60

# A clause boundary is a place a title can be cut and still read as a finished
# phrase. Below _CLAUSE_FLOOR characters the remainder is too short to describe
# the page, so a word-boundary trim of the full headline beats it.
_CLAUSE_RE = re.compile(r'[,:;]|\s[-–—]\s')
_CLAUSE_FLOOR = 30

# Words that open a phrase rather than close one. A title ending on one reads as
# cut off ("... Section 63 Sources, Conditions &").
_DANGLING = ('and', 'the', 'a', 'an', 'of', 'for', 'in', 'on', 'to', 'with', '&')


def shorten_title(text, maxlen=TITLE_MAX):
    """Cut a long headline down to the width Google displays, ending on a
    complete phrase wherever possible.

    Prefers the last clause boundary — comma, colon, or spaced dash — because
    cutting there finishes a thought. Falls back to a word-boundary trim, which
    can strand the first word of a phrase ("... ADT-1, Sections"); nothing in the
    string distinguishes a dangling 'Raise' from a fine 'Removal', so the pages
    that matter carry a hand-written title in SEO_TITLES instead."""
    text = ' '.join(str(text or '').split())
    if len(text) <= maxlen:
        return text

    boundary = None
    for m in _CLAUSE_RE.finditer(text):
        if _CLAUSE_FLOOR <= m.start() <= maxlen:
            boundary = m.start()
    if boundary is not None:
        return text[:boundary].rstrip(' ,;:-–—&')

    cut = text[:maxlen]
    if ' ' in cut:
        cut = cut[:cut.rindex(' ')]
    cut = cut.rstrip(' ,;:-–—&')
    while True:
        head, _, last = cut.rpartition(' ')
        if head and last.lower() in _DANGLING:
            cut = head.rstrip(' ,;:-–—&')
            continue
        return cut


@app.template_filter('fit')
def fit(base, *extras, maxlen=TITLE_MAX):
    """Build a <title> from a base plus optional suffixes, keeping the whole
    thing inside the width Google displays.

    Each extra is appended only if it still fits; one that doesn't is skipped
    rather than ending the chain, so a short brand can still land when a long
    value-proposition suffix cannot. Order them most valuable first."""
    out = shorten_title(base, maxlen)
    for extra in extras:
        if len(out) + len(extra) <= maxlen:
            out += extra
    return out


@app.template_filter('seotitle')
def seotitle(article, brand=' - Law Minded', maxlen=TITLE_MAX):
    """Build a search-friendly <title>. Prefers a hand-written SEO_TITLES entry,
    then the article's admin-editable seo_title column, and otherwise shortens
    the long editorial headline. The ' - Law Minded' suffix is only appended when
    the whole thing still fits Google's ~60-char display width, so titles stop
    truncating mid-phrase in search results."""
    def _get(row, key):
        try:
            v = row[key]
        except (KeyError, IndexError, TypeError):
            v = None
        return v.strip() if isinstance(v, str) else ''

    # Three sources, most deliberate first: a title written for the query, the
    # admin-editable column, then the shortened headline.
    title = SEO_TITLES.get(_get(article, 'slug'), '') or _get(article, 'seo_title')
    if not title:
        # Do NOT reduce the title to the text before the colon. That rule used
        # to fire on 113 of 126 articles, and on this site the part after the
        # colon is where the statute reference lives — "Appointment of KMP:
        # Section 203 Thresholds" rendered as just "Appointment of KMP".
        # Search Console for Aug 2026 shows the section numbers are precisely
        # what people search ("section 203 of companies act 2013", "section 68",
        # "62(1)(a)"), so dropping them removed the strongest on-page signal
        # from the one tag that carries the most weight. shorten_title keeps as
        # much of the reference as fits, cutting at a clause boundary so the
        # result still reads as a finished phrase.
        title = shorten_title(_get(article, 'title'), maxlen)
    if brand and len(title) + len(brand) <= maxlen:
        return title + brand
    return title


@app.template_filter('metadesc')
def metadesc(value, maxlen=155):
    """Trim a description to the width Google and Bing actually render in a
    search snippet (~155 chars). Anything longer is cut off mid-word by the
    search engine itself, so we cut it first — on a word boundary, dropping any
    dangling punctuation — and the snippet reads as a finished phrase. Applied
    once in base.html, so it covers every page's description, og:description
    and twitter:description."""
    s = ' '.join(str(value or '').split())
    if len(s) <= maxlen:
        return s
    cut = s[:maxlen]
    if ' ' in cut:
        cut = cut[:cut.rindex(' ')]
    return cut.rstrip(' ,;:-–—')


# One alternation over every linkable phrase, longest first so "GST registration"
# wins over "GST". The lookarounds do the job of \b but also refuse to fire mid-URL
# or inside a hyphenated token, so "section 185" never matches "sub-section 185a".
_LINK_RE = re.compile(
    r'(?<![\w\-/])(' + '|'.join(
        re.escape(p) for p in sorted(INTERNAL_LINKS, key=len, reverse=True)
    ) + r')(?![\w\-])', re.I)

# Headings stay unlinked (they are the page's own outline, not a route out), and
# nesting an <a> inside an <a> is invalid HTML that browsers silently mangle.
_NO_LINK_ZONE = re.compile(r'</?(a|h[1-6])\b', re.I)


@app.template_filter('autolink')
def autolink(content, current_slug='', limit=8):
    """Turn key phrases in an article body into links to the article that
    explains them.

    Editors write the guides as standalone pieces, so nothing cross-references
    anything — 123 articles between them held 7 links, all pointing off-site.
    That leaves every guide an island: readers hit a dead end, and crawlers get
    no path between pages or any anchor text describing them.

    Linking here rather than in the stored HTML means new articles are covered
    the moment they publish, and a slug rename never leaves a dead link behind.
    Only the first mention of each target is linked, capped at `limit` per page,
    so the result reads like an editor added it rather than a machine."""
    if not content:
        return content
    # Targets an editor already linked by hand count against the budget, so a
    # manually-linked article never picks up a second link to the same page.
    used = set(re.findall(r'href="/article/([a-z0-9\-]+)"', content))

    def _link_text(text):
        def repl(m):
            phrase = m.group(0)
            slug = INTERNAL_LINKS.get(phrase.lower())
            if not slug or slug == current_slug or slug in used or len(used) >= limit:
                return phrase
            used.add(slug)
            return f'<a href="{url_for("article", slug=slug)}">{phrase}</a>'
        return _LINK_RE.sub(repl, text)

    out, pos, depth = [], 0, 0
    for m in re.finditer(r'<[^>]+>', content):
        chunk = content[pos:m.start()]
        out.append(_link_text(chunk) if depth == 0 else chunk)
        tag = m.group(0)
        out.append(tag)
        if _NO_LINK_ZONE.match(tag):
            depth += 1 if tag[1] != '/' else -1
            depth = max(depth, 0)
        pos = m.end()
    tail = content[pos:]
    out.append(_link_text(tail) if depth == 0 else tail)
    return ''.join(out)


@app.template_filter('faqs')
def faqs(content, limit=10):
    """Pull (question, answer) pairs from an article's 'Frequently asked
    questions' section for FAQPage schema. Matches the seeded structure:
    an <h2>FAQ</h2> heading followed by <p><strong>Q?</strong> A</p> pairs.
    ponytail: naive regex parse of known-good seeded HTML; re-check if the
    editor ever emits a different FAQ structure."""
    if not content:
        return []
    import re
    m = re.search(
        r'<h2[^>]*>\s*(?:frequently asked questions|common questions|faqs?)\s*</h2>(.*?)(?:<h2|$)',
        content, re.I | re.S)
    if not m:
        return []
    out = []
    for q, a in re.findall(r'<p>\s*<strong>(.*?)</strong>\s*(.*?)</p>', m.group(1), re.S):
        q = re.sub(r'<[^>]+>', '', q).strip()
        a = re.sub(r'<[^>]+>', '', a).strip()
        if q.endswith('?') and a:
            out.append({'q': q, 'a': a})
    return out[:limit]


def _article_image_url(slug):
    """Static URL of an article's header image if the file exists, else None.
    Used for the article hero, og:image, and the blog-listing card thumbnails."""
    if slug and os.path.exists(os.path.join(app.static_folder, 'img', 'articles', f'{slug}.webp')):
        return url_for('static', filename=f'img/articles/{slug}.webp')
    return None


app.jinja_env.globals['article_image'] = _article_image_url
app.jinja_env.globals['seo_descriptions'] = SEO_DESCRIPTIONS


@app.context_processor
def inject_globals():
    return {
        'adsense_client': ADSENSE_CLIENT if ADS_ENABLED else '',
        'adsense_slots': ADSENSE_SLOTS,
        'google_site_verification': GOOGLE_SITE_VERIFICATION,
        'category_map': CATEGORY_MAP,
        'topics': C.TOPICS,
        'current_year': date.today().year,
        'site_url': SITE_URL,
        'canonical_url': SITE_URL + request.path,
        'asset_v': asset_version,
        'author': AUTHOR,
    }


# Both domains serve directly. There is deliberately no host redirect here:
# lawminded.in and lawminded.co.in each answer 200 on every path, as do their www
# forms, because the owner wants both reachable rather than one funnelling into
# the other. nginx still sends http -> https, but it preserves $host, so it never
# moves a visitor between domains.
#
# What stops that becoming a duplicate-content problem is the canonical tag, not
# a redirect. `canonical_url` above is built from SITE_URL, so every page served
# on any hostname declares the SITE_URL copy as the one to index, and the ranking
# signals still consolidate onto a single address. Sitemap, OG and JSON-LD URLs
# come from the same constant and agree with it.
#
# If you ever want one domain to funnel into the other again, this is the place:
# reinstate a before_request that 301s to the host in SITE_URL, and let
# /.well-known/ through untouched so certbot renewal keeps working.


# ─── Helpers ────────────────────────────────────────────────────────────────

def slugify(text):
    text = text.lower().strip()
    text = re.sub(r'[^\w\s-]', '', text)
    return re.sub(r'[\s_-]+', '-', text)


def get_articles_by_cat():
    db = get_db()
    by_cat = {}
    for cat in CATEGORY_MAP:
        by_cat[cat] = db.execute(
            'SELECT * FROM articles WHERE category=? AND published=1 ORDER BY created_at DESC',
            (cat,)
        ).fetchall()
    db.close()
    return by_cat


def get_all_articles():
    db = get_db()
    rows = db.execute(
        'SELECT * FROM articles WHERE published=1 ORDER BY created_at DESC'
    ).fetchall()
    db.close()
    return rows


def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('admin_logged_in'):
            return redirect(url_for('admin_login'))
        return f(*args, **kwargs)
    return decorated


# ─── Documents (DB-managed templates & resolutions) ──────────────────────────
DOC_TYPES = {'template': 'Template', 'board': 'Board Resolution'}
DOC_LIST_TITLE = {'board': 'Board Resolutions'}


def get_documents(doc_type):
    db = get_db()
    rows = db.execute(
        'SELECT * FROM documents WHERE doc_type=? ORDER BY sort_order, id', (doc_type,)
    ).fetchall()
    db.close()
    return rows


def get_document(doc_type, slug):
    db = get_db()
    row = db.execute(
        'SELECT * FROM documents WHERE doc_type=? AND slug=?', (doc_type, slug)
    ).fetchone()
    db.close()
    return row


# ─── Document Formats (uploaded Word files, DB-managed) ──────────────────────
FORMATS_DIR = os.path.join(app.static_folder, 'formats')


def get_formats_grouped():
    """Shape DB format rows into the {name, icon, docs:[{slug,file,title,desc}]}
    structure the Templates page expects, preserving the known category order."""
    db = get_db()
    rows = db.execute('SELECT * FROM formats ORDER BY sort_order, id').fetchall()
    db.close()
    by_cat = {}
    for r in rows:
        by_cat.setdefault(r['category'], []).append(
            {'slug': r['slug'], 'file': r['filename'], 'title': r['title'],
             'desc': r['description'] or ''}
        )
    cats, seen = [], set()
    for name in F.CATEGORY_NAMES:
        if name in by_cat:
            cats.append({'name': name, 'icon': F.CATEGORY_ICONS.get(name, F.DEFAULT_CATEGORY_ICON),
                         'docs': by_cat[name]})
            seen.add(name)
    for name, docs in by_cat.items():
        if name not in seen:
            cats.append({'name': name, 'icon': F.CATEGORY_ICONS.get(name, F.DEFAULT_CATEGORY_ICON),
                         'docs': docs})
    return cats


def get_format(slug):
    db = get_db()
    row = db.execute('SELECT * FROM formats WHERE slug=?', (slug,)).fetchone()
    db.close()
    return row


def formats_count():
    db = get_db()
    n = db.execute('SELECT COUNT(*) FROM formats').fetchone()[0]
    db.close()
    return n


def doc_to_view(d):
    """Shape a DB document row for the public templates/resolutions pages."""
    return dict(
        d,
        desc=d['description'],
        tags=(d['tags'].split(',') if d['tags'] else []),
        html=C.render_resolution_html(C.parse_doc_body(d['body'])),
    )


def get_setting(key):
    db = get_db()
    row = db.execute('SELECT value FROM settings WHERE key=?', (key,)).fetchone()
    db.close()
    return row['value'] if row else None


def set_setting(key, value):
    db = get_db()
    db.execute(
        'INSERT INTO settings (key, value) VALUES (?,?) '
        'ON CONFLICT(key) DO UPDATE SET value=excluded.value',
        (key, value)
    )
    db.commit()
    db.close()


def _set_env_vars(updates):
    """Persist key=value pairs to .env (atomic) and the live process env."""
    lines = []
    if os.path.exists(ENV_PATH):
        with open(ENV_PATH, 'r') as f:
            lines = f.read().splitlines()
    seen, out = set(), []
    for ln in lines:
        key = ln.split('=', 1)[0] if ('=' in ln and not ln.lstrip().startswith('#')) else None
        if key in updates:
            out.append(f'{key}={updates[key]}'); seen.add(key)
        else:
            out.append(ln)
    for k, v in updates.items():
        if k not in seen:
            out.append(f'{k}={v}')
    tmp = ENV_PATH + '.tmp'
    with open(tmp, 'w') as f:
        f.write('\n'.join(out) + '\n')
    os.replace(tmp, ENV_PATH)
    for k, v in updates.items():
        os.environ[k] = v


def _env_value(key):
    """Read a key straight from .env each time, so every gunicorn worker sees a
    credential change immediately (no restart). Falls back to the process env."""
    try:
        with open(ENV_PATH, 'r') as f:
            for ln in f:
                ln = ln.rstrip('\n')
                if ln.startswith(key + '='):
                    return ln.split('=', 1)[1]
    except OSError:
        pass
    return os.getenv(key, '')


def get_admin_username():
    return (_env_value('ADMIN_USERNAME') or '').strip()


def _admin_hash():
    """Admin password hash, stored base64-encoded in .env (keeps the '$'-laden
    hash safe from dotenv variable interpolation)."""
    b64 = _env_value('ADMIN_PW_HASH_B64') or ''
    try:
        return base64.b64decode(b64).decode('utf-8') if b64 else ''
    except Exception:
        return ''


def admin_configured():
    return bool(get_admin_username() and _admin_hash())


def check_admin(username, password):
    """Verify admin username + password. Fails closed — there is no default."""
    u, h = get_admin_username(), _admin_hash()
    if not u or not h:
        return False
    if not hmac.compare_digest((username or '').strip().lower(), u.lower()):
        return False
    return check_password_hash(h, password or '')


def set_admin_credentials(username=None, password=None):
    """Write new admin credentials to .env (password stored only as a hash)."""
    updates = {}
    if username is not None:
        updates['ADMIN_USERNAME'] = username.strip()
    if password is not None:
        h = generate_password_hash(password, method='pbkdf2:sha256')
        updates['ADMIN_PW_HASH_B64'] = base64.b64encode(h.encode('utf-8')).decode('ascii')
    if updates:
        _set_env_vars(updates)


def build_resolution_docx(title, blocks):
    """Build a Word (.docx) document from resolution blocks, returned as BytesIO."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for kind, text in blocks:
        if kind == 'spacer':
            doc.add_paragraph('')
        elif kind in ('heading', 'subheading'):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            if kind == 'heading':
                run.font.size = Pt(12)
        elif kind == 'bullet':
            doc.add_paragraph(text, style='List Bullet')
        else:
            doc.add_paragraph(text)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ─── Public Pages ─────────────────────────────────────────────────────────────

@app.route('/')
def index():
    articles = get_all_articles()
    # Counted live, never hand-typed. The homepage used to claim "50+ topics"
    # and "10+ templates" against a real 131 and 55, and the numbers only
    # appeared after JS ran — so crawlers that don't execute JS read "0+".
    stats = {
        'articles': len(articles),
        'formats': formats_count(),
        'topics': len(C.TOPICS),
        'judgments': len(C.JUDGMENTS),
    }
    return render_template('index.html', featured=articles[:6], stats=stats,
                           hero_image='/static/img/pages/home.webp')


@app.route('/about')
def about():
    return render_template('about.html')


@app.route('/author/<slug>')
def author_page(slug):
    """Retired at the owner's request — the author bio page no longer exists.

    Kept as a 301 rather than a 404 because the URL was in the sitemap and may
    be indexed; /about is the nearest surviving page about who publishes this
    site. The byline still names the author and the credential, it just does
    not link anywhere.
    """
    return redirect(url_for('about'), code=301)


@app.route('/blogs')
def blogs():
    return render_template('blogs.html', articles_by_cat=get_articles_by_cat())


@app.route('/topic/<slug>')
def topic(slug):
    """Landing page for one category — the hub half of the hub-and-spoke.

    /blogs lists all 123 guides on a single page, so no URL ever targeted a
    topic as a phrase and the guides within a topic had no shared parent. This
    gives each category a page of its own that links to every article under it;
    the article breadcrumb links back."""
    cat = C.TOPIC_BY_SLUG.get(slug)
    if not cat:
        abort(404)
    db = get_db()
    articles = db.execute(
        'SELECT * FROM articles WHERE category=? AND published=1 '
        'ORDER BY created_at DESC', (cat,)
    ).fetchall()
    db.close()
    return render_template('topic.html', topic=C.TOPICS[cat], cat=cat,
                           articles=articles)


@app.route('/article/<slug>')
def article(slug):
    # A retired duplicate keeps its URL working: 301 to the guide that replaced
    # it so any existing link or index entry lands on the survivor rather than
    # a 404. Checked before the query — the row is still there, just unpublished.
    if slug in RETIRED_ARTICLES:
        return redirect(url_for('article', slug=RETIRED_ARTICLES[slug]), code=301)
    db = get_db()
    row = db.execute(
        'SELECT * FROM articles WHERE slug=? AND published=1', (slug,)
    ).fetchone()
    if not row:
        db.close()
        abort(404)
    related = db.execute(
        'SELECT * FROM articles WHERE category=? AND id!=? AND published=1 ORDER BY created_at DESC LIMIT 4',
        (row['category'], row['id'])
    ).fetchall()
    # Top up to 4 with recent articles from other categories so the row is always full.
    if len(related) < 4:
        seen = {r['id'] for r in related} | {row['id']}
        extra = db.execute(
            'SELECT * FROM articles WHERE published=1 ORDER BY created_at DESC LIMIT 12'
        ).fetchall()
        related = list(related) + [r for r in extra if r['id'] not in seen][:4 - len(related)]
    db.close()
    return render_template('article.html', article=row, related=related,
                           hero_image=_article_image_url(slug))


# ─── Weekly draft: preview on a phone, publish with one tap ──────────────────
# The scheduled writer (see automation/) inserts its article with
# published=0, so /article already 404s it, then sends a signed link here. The
# signature is the only credential: a leaked link can publish a draft we wrote
# ourselves, which is a smaller risk than making the owner log into /admin on a
# phone at 7am. Approval is a POST — Telegram fetches links to build previews,
# and a GET that publishes would fire the moment the message rendered.
def _load_draft(slug, token):
    try:
        if _draft_serializer.loads(token or '') != slug:
            abort(404)
    except BadSignature:
        abort(404)
    db = get_db()
    row = db.execute(
        'SELECT * FROM articles WHERE slug=? AND published=0', (slug,)
    ).fetchone()
    db.close()
    if not row:
        abort(404)
    return row


@app.route('/draft/<slug>')
def draft_preview(slug):
    token = request.args.get('t')
    row = _load_draft(slug, token)
    return render_template('article.html', article=row, related=[],
                           hero_image=_article_image_url(slug),
                           draft_token=token)


@app.route('/draft/<slug>/publish', methods=['POST'])
def draft_publish(slug):
    row = _load_draft(slug, request.form.get('t'))
    db = get_db()
    # Stamp the publication date at approval time, not at draft time, so the
    # date on the page is the date it actually went live — in IST, since that is
    # what every other date in this table is and what the reader expects.
    now = ist_now()
    db.execute('UPDATE articles SET published=1, publish_on=NULL, created_at=?, '
               'updated_at=? WHERE id=?', (now, now, row['id']))
    db.commit()
    db.close()
    # Subscribers hear about it the moment it goes live. Backgrounded so the tap
    # returns straight to the article rather than waiting on a mail server.
    announce_article_async(slug)
    flash('Published. Subscribers are being notified.', 'success')
    return redirect(url_for('article', slug=slug))


@app.route('/draft/<slug>/reject', methods=['POST'])
def draft_reject(slug):
    row = _load_draft(slug, request.form.get('t'))
    db = get_db()
    db.execute('DELETE FROM articles WHERE id=?', (row['id'],))
    db.commit()
    db.close()
    flash('Draft rejected and removed.', 'success')
    return redirect(url_for('index'))


@app.route('/templates')
def templates_page():
    # Old starter-draft templates retired; the real Word Formats Library is now
    # the sole template offering on this page.
    return render_template('templates_page.html',
                           format_categories=get_formats_grouped(),
                           formats_count=formats_count())


@app.route('/template/<slug>/download')
def template_download(slug):
    item = get_document('template', slug)
    if not item:
        abort(404)
    bio = build_resolution_docx(item['title'], C.parse_doc_body(item['body']))
    return send_file(
        bio, as_attachment=True, download_name=f'{slug}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/format/<slug>/download')
def format_download(slug):
    """Serve one of the real .docx document formats from static/formats.

    The slug is looked up in FORMATS_BY_SLUG, so only known filenames are ever
    served (no path-traversal from user input)."""
    item = get_format(slug)
    if not item:
        abort(404)
    folder = os.path.join(app.static_folder, 'formats')
    return send_from_directory(
        folder, item['filename'], as_attachment=True, download_name=item['filename'],
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@lru_cache(maxsize=128)
def _render_format_preview(slug):
    """Render a stored .docx format to safe HTML for the in-page preview modal.

    Returns an HTML fragment (headings, paragraphs, tables) built from our own
    trusted files, with all text escaped. Cached per slug so repeat previews
    don't re-parse the document."""
    item = get_format(slug)
    if not item:
        return None
    from markupsafe import escape
    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table as _Tbl
    from docx.text.paragraph import Paragraph as _Para

    doc = Document(os.path.join(app.static_folder, 'formats', item['filename']))

    def runs_html(p):
        parts = []
        for r in p.runs:
            t = str(escape(r.text or ''))
            if not t:
                continue
            if r.bold:
                t = f'<strong>{t}</strong>'
            if r.italic:
                t = f'<em>{t}</em>'
            parts.append(t)
        return ''.join(parts) if parts else str(escape(p.text or ''))

    def cell_html(cell):
        bits = [runs_html(p) for p in cell.paragraphs if p.text.strip()]
        return '<br>'.join(bits)

    out = []
    for child in doc.element.body.iterchildren():
        if child.tag == qn('w:p'):
            p = _Para(child, doc)
            txt = runs_html(p)
            if not txt.strip():
                continue
            style = (p.style.name if p.style else '') or ''
            tag = 'h4' if (style.startswith('Heading') or style.startswith('Title')) else 'p'
            out.append(f'<{tag}>{txt}</{tag}>')
        elif child.tag == qn('w:tbl'):
            tbl = _Tbl(child, doc)
            rows = ''.join(
                '<tr>' + ''.join(f'<td>{cell_html(c)}</td>' for c in row.cells) + '</tr>'
                for row in tbl.rows
            )
            if rows:
                out.append(f'<div class="preview-table-wrap"><table class="preview-table"><tbody>{rows}</tbody></table></div>')
    return '\n'.join(out)


@app.route('/format/<slug>/preview')
def format_preview(slug):
    html_out = _render_format_preview(slug)
    if html_out is None:
        abort(404)
    return html_out


@app.route('/format/<slug>')
def format_page(slug):
    """A crawlable page per document.

    All 55 formats used to live behind buttons on the single /templates URL, so
    nothing could rank for "board resolution format" or "share transfer deed" —
    the queries people actually type. Each document now has its own URL with the
    full text rendered into the page, not just a download link.
    """
    item = get_format(slug)
    if not item:
        abort(404)
    db = get_db()
    related = db.execute(
        'SELECT slug, title, description FROM formats WHERE category=? AND slug!=? '
        'ORDER BY sort_order, id LIMIT 6', (item['category'], slug)
    ).fetchall()
    total = db.execute('SELECT COUNT(*) FROM formats').fetchone()[0]
    db.close()
    return render_template('format.html', fmt=item, related=related,
                           formats_total=total,
                           preview=_render_format_preview(slug),
                           hero_image='/static/img/pages/templates.webp')


@app.route('/compare')
def compare():
    """Index of every comparison, with the first one rendered in full.

    The table markup is emitted server-side. Until Aug 2026 this page shipped
    an empty <div> and let JS inject the table, so search engines and AI
    crawlers saw five headings and no data.
    """
    return render_template('compare.html',
                           comparisons=C.COMPARISON_TABLES,
                           current=None,
                           hero_image='/static/img/pages/compare.webp')


@app.route('/compare/<slug>')
def compare_one(slug):
    current = C.COMPARISON_BY_SLUG.get(slug)
    if current is None:
        abort(404)
    return render_template('compare.html',
                           comparisons=C.COMPARISON_TABLES,
                           current=current,
                           hero_image='/static/img/pages/compare.webp')


@app.route('/judgments')
def judgments():
    return render_template('judgments.html', judgments=C.JUDGMENTS)


@app.route('/judgment/<slug>')
def judgment(slug):
    j = next((x for x in C.JUDGMENTS if x['slug'] == slug), None)
    if not j:
        abort(404)
    return render_template('judgment.html', j=j, published=JUDGMENTS_PUBLISHED)


@app.route('/resolutions/<rtype>')
def resolutions(rtype):
    if rtype not in DOC_LIST_TITLE:
        abort(404)
    rendered = [doc_to_view(d) for d in get_documents(rtype)]
    return render_template('resolutions.html', rtype=rtype,
                           list_title=DOC_LIST_TITLE[rtype], items=rendered)


@app.route('/resolution/<rtype>/<slug>/download')
def resolution_download(rtype, slug):
    if rtype not in DOC_LIST_TITLE:
        abort(404)
    item = get_document(rtype, slug)
    if not item:
        abort(404)
    bio = build_resolution_docx(item['title'], C.parse_doc_body(item['body']))
    return send_file(
        bio, as_attachment=True, download_name=f'{slug}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/resources')
def resources():
    return render_template('resources.html', resources=C.RESOURCES)


@app.route('/faq')
def faq():
    return render_template('faq.html', faqs=C.FAQS)


@app.route('/contact')
def contact_page():
    return render_template('contact.html')


@app.route('/terms')
def terms():
    return render_template('terms.html')


@app.route('/privacy')
def privacy():
    return render_template('privacy.html')


@app.route('/search')
def search():
    query = request.args.get('q', '').strip()
    results = C.search_all(query, get_all_articles()) if query else []
    if query:
        q = query.lower()
        db = get_db()
        docs = db.execute('SELECT * FROM documents').fetchall()
        db.close()
        for d in docs:
            hay = ' '.join(str(d[k] or '') for k in ('title', 'description', 'tags', 'body')).lower()
            if q in hay:
                if d['doc_type'] == 'template':
                    continue  # retired starter templates — no longer surfaced
                if d['doc_type'] in DOC_LIST_TITLE:
                    results.append({'type': DOC_TYPES.get(d['doc_type'], 'Document'), 'title': d['title'],
                                    'snippet': d['description'], 'url_kind': 'resolutions', 'url_arg': d['doc_type']})
    return render_template('search.html', query=query, results=results)


# ─── Form / API Endpoints ─────────────────────────────────────────────────────

@app.route('/contact', methods=['POST'])
@limiter.limit('3 per hour; 10 per day')
def contact():
    # Every rejection below returns the same success payload a real submission
    # gets. Telling a bot which check caught it just helps it adapt.
    ok = jsonify({'success': True, 'message': 'Your query has been submitted successfully!'})

    # Honeypot: real users never fill the hidden "website" field; bots do.
    if request.form.get('website', '').strip():
        return ok

    # Dwell time: the form must have been rendered a few seconds ago. Scripted
    # posts either omit the token or replay one far too fast.
    if not form_token_ok(request.form.get('ft', '')):
        return ok

    name = request.form.get('name', '').strip()
    email = request.form.get('email', '').strip()
    query = request.form.get('query', '').strip()

    if looks_like_spam(name, email, query):
        return ok

    errors = {}
    if not name:
        errors['name'] = 'Name is required.'
    if not email or not VALID_EMAIL_RE.match(email):
        errors['email'] = 'Valid email is required.'
    if not query or len(query) < 20:
        errors['query'] = 'Please describe your query (at least 20 characters).'

    if errors:
        return jsonify({'success': False, 'errors': errors}), 400

    db = get_db()
    db.execute(
        'INSERT INTO contact_messages (name, email, query) VALUES (?,?,?)',
        (name, email, query)
    )
    db.commit()
    db.close()

    try:
        # Internal notification to the team (reply goes straight to the sender).
        notify = Message(
            subject=f'New enquiry from {name} — Law Minded',
            recipients=[CONTACT_RECEIVER],
            body=f'Name: {name}\nEmail: {email}\n\nMessage:\n{query}'
        )
        notify.reply_to = email
        mail.send(notify)
        # Branded acknowledgement to the person who wrote in.
        send_branded_email(
            subject='We received your message — Law Minded',
            recipients=[email],
            heading=f'Thanks, {name} — we have your message',
            body_html=(
                '<p style="margin:0 0 12px;font:400 15px/1.65 Arial,sans-serif;">'
                'Thanks for reaching out to Law Minded. We have received your message '
                'and will get back to you as soon as we can.</p>'
                '<p style="margin:0;font:400 15px/1.65 Arial,sans-serif;color:#6F6857;">'
                'You can simply reply to this email if you have anything to add.</p>'
            ),
            body_text=(
                f'Hi {name},\n\n'
                'Thanks for reaching out to Law Minded. We have received your message '
                'and will get back to you as soon as we can.\n\n'
                '— Law Minded'
            ),
        )
    except Exception:
        pass

    return jsonify({'success': True, 'message': 'Your query has been submitted successfully!'})


@app.route('/newsletter', methods=['POST'])
@limiter.limit('6 per minute')
def newsletter():
    if request.form.get('website', '').strip():
        return jsonify({'success': True, 'message': "Thank you! You've been subscribed."})

    name = request.form.get('name', '').strip()
    email = request.form.get('email', '').strip()

    if not email or not VALID_EMAIL_RE.match(email):
        return jsonify({'success': False, 'message': 'Please enter a valid email address.'}), 400

    db = get_db()
    try:
        db.execute('INSERT INTO subscribers (name, email) VALUES (?,?)', (name, email))
        db.commit()
    except sqlite3.IntegrityError:
        db.close()
        return jsonify({'success': True, 'message': "You're already subscribed!"})
    db.close()

    try:
        unsub = unsubscribe_url(email)
        hi = name or 'there'
        hub = f'{SITE_URL}{url_for("blogs")}'
        send_branded_email(
            subject='Welcome to Law Minded',
            recipients=[email],
            heading=f'Welcome, {hi}',
            body_html=(
                '<p style="margin:0 0 12px;font:400 15px/1.65 Arial,sans-serif;">'
                'You are now subscribed — thank you for joining.</p>'
                '<p style="margin:0 0 12px;font:400 15px/1.65 Arial,sans-serif;">'
                'From time to time we will send plain-English updates on compliance changes, '
                'new guides, and important court judgments. Nothing more, and nothing you did not sign up for.</p>'
                '<p style="margin:0;font:400 15px/1.65 Arial,sans-serif;">'
                f'In the meantime, browse our free <a href="{hub}" style="color:#8A5E07;">Knowledge Hub</a> '
                'for guides written for founders, professionals, and everyday citizens.</p>'
            ),
            body_text=(
                f'Hi {hi},\n\n'
                'You are now subscribed to Law Minded — thank you for joining.\n\n'
                'From time to time we will send plain-English updates on compliance changes, '
                'new guides, and important court judgments.\n\n'
                f'Browse our free Knowledge Hub: {hub}\n\n'
                '— Law Minded\n\n'
                f'Unsubscribe anytime: {unsub}'
            ),
            unsub=unsub,
        )
    except Exception:
        pass

    return jsonify({'success': True, 'message': f'Thank you{", " + name if name else ""}! You\'ve been subscribed.'})


@app.route('/unsubscribe', methods=['GET', 'POST'])
@csrf.exempt
def unsubscribe():
    """Unsubscribe via a signed token.

    GET only shows a confirmation page with an "Unsubscribe" button — visiting the
    link never removes anyone. The removal happens on POST: either the in-page
    button, or the RFC 8058 one-click path that Gmail/Outlook fire automatically
    from the List-Unsubscribe header (identified by the List-Unsubscribe=One-Click
    body field, which the button submit does not send)."""
    token = request.values.get('token', '')
    try:
        email = _unsub_serializer.loads(token)
    except BadSignature:
        email = None

    # A plain visit changes nothing — just ask the person to confirm.
    if request.method == 'GET':
        return render_template('unsubscribe.html', email=email, token=token,
                               removed=False, confirm=bool(email), valid=bool(email))

    # POST: actually remove them.
    removed = False
    if email:
        db = get_db()
        cur = db.execute('DELETE FROM subscribers WHERE email=?', (email,))
        db.commit()
        removed = cur.rowcount > 0
        db.close()

    # Mail-client one-click (RFC 8058) expects a bare 2xx with no page body.
    if request.form.get('List-Unsubscribe') == 'One-Click':
        return ('', 204)

    return render_template('unsubscribe.html', email=email, token=token,
                           removed=removed, confirm=False, valid=bool(email))


@app.route('/download-request', methods=['POST'])
@limiter.limit('20 per minute')
def download_request():
    email = request.form.get('email', '').strip()
    template_name = request.form.get('template', '').strip()

    if email and VALID_EMAIL_RE.match(email):
        db = get_db()
        db.execute(
            'INSERT INTO download_requests (email, template_name) VALUES (?,?)',
            (email, template_name)
        )
        try:
            db.execute('INSERT INTO subscribers (name, email) VALUES (?,?)', ('', email))
        except sqlite3.IntegrityError:
            pass
        db.commit()
        db.close()

    return jsonify({'success': True, 'message': 'Download ready!'})


@app.route('/ads.txt')
def ads_txt():
    # Required by Google AdSense to authorise this site to show your ads.
    if not ADSENSE_CLIENT:
        abort(404)
    pub = ADSENSE_CLIENT.replace('ca-', '')
    line = f'google.com, {pub}, DIRECT, f08c47fec0942fa0\n'
    return app.response_class(line, mimetype='text/plain')


# Google Search Console ownership proof. The path and the body both have to be
# exactly what Google issued — it fetches the file and string-matches it — so
# this returns the line verbatim, with no template, no trailing newline and no
# styling. Both domains answer it, so this one token verifies whichever property
# gets registered. If a second property is added it issues its own token, which
# means another copy of this route with the new name; deliberately not a
# wildcard, because a catch-all on /<anything>.html at the site root is a wide
# rule to leave lying around for one static string.
@app.route('/google790ff7b7719dd579.html')
def google_site_verification():
    return app.response_class(
        'google-site-verification: google790ff7b7719dd579.html',
        mimetype='text/html')


@app.route('/robots.txt')
def robots_txt():
    lines = [
        'User-agent: *',
        'Allow: /',
        'Disallow: /admin',
        'Disallow: /admin/',
        f'Sitemap: {SITE_URL}/sitemap.xml',
        '',
    ]
    return app.response_class('\n'.join(lines), mimetype='text/plain')


@app.route('/sitemap')
def sitemap_page():
    # Human-readable HTML site map (this is what the footer "Sitemap" link opens).
    # Distinct from the machine-readable /sitemap.xml below, which is for search engines.
    return render_template('sitemap.html', judgments=C.JUDGMENTS)


@app.route('/sitemap.xml')
def sitemap_xml():
    """Machine-readable sitemap, generated LIVE on every request.

    Auto-updating: every published article is pulled from the DB below, so a new
    blog (added via blog_seed*.py or the admin) appears here automatically after
    deploy — no manual editing of this list is ever needed. A brand-new *category*
    only needs adding to CATEGORY_MAP in content.py (it rides on the /blogs URL).
    robots.txt advertises this file; the human-facing index is /sitemap."""
    def freshest(*dates):
        """The latest of the dates given, ignoring blanks. ISO dates sort
        lexicographically, so max() is the whole comparison."""
        return max((d[:10] for d in dates if d), default=None)

    # Read the article rows once: they date the articles, and the topic hub and
    # index pages are dated by the newest article they list.
    db = get_db()
    article_rows = db.execute(
        'SELECT slug, category, updated_at FROM articles WHERE published=1 '
        'ORDER BY updated_at DESC').fetchall()
    format_rows = db.execute(
        'SELECT slug, updated_at FROM formats ORDER BY sort_order, id').fetchall()
    db.close()

    newest_article = freshest(*(r['updated_at'] or '' for r in article_rows))
    newest_format = freshest(*(r['updated_at'] or '' for r in format_rows))
    by_category = {}
    for r in article_rows:
        by_category.setdefault(r['category'], []).append(r['updated_at'] or '')

    # Static pages. A lastmod is given only where a real date backs it: the
    # listing pages are dated by what they list, and the two whose own markup
    # changed in the metadata revision carry that date. The rest — terms,
    # privacy, contact — have no date we can honestly state, and an invented one
    # is worse than none.
    pages = [
        ('index', 'weekly', '1.0'),
        ('about', 'monthly', '0.7'),
        ('blogs', 'daily', '0.9'),
        ('templates_page', 'weekly', '0.8'),
        ('compare', 'monthly', '0.7'),
        ('judgments', 'monthly', '0.7'),
        ('resources', 'monthly', '0.6'),
        ('faq', 'monthly', '0.7'),
        ('contact_page', 'yearly', '0.5'),
        ('terms', 'yearly', '0.3'),
        ('privacy', 'yearly', '0.3'),
        ('sitemap_page', 'yearly', '0.3'),
    ]
    static_lastmod = {
        'index': freshest(newest_article, SEARCH_META_CHANGED),
        'blogs': newest_article,
        'templates_page': freshest(newest_format, SEARCH_META_CHANGED),
        'judgments': JUDGMENTS_PUBLISHED,
    }
    urls = []
    for endpoint, freq, prio in pages:
        urls.append((SITE_URL + url_for(endpoint), freq, prio,
                     static_lastmod.get(endpoint)))

    # Topic hub pages — one per category, added automatically with the category.
    # A hub lists its category's articles, so it genuinely changes whenever the
    # newest of them does.
    for cat, t in C.TOPICS.items():
        urls.append((SITE_URL + url_for('topic', slug=t['slug']), 'weekly', '0.8',
                     freshest(*by_category.get(cat, []))))

    # One URL per act comparison — these are the pages that can actually rank
    # for "x vs y" queries; /compare alone never could.
    for c in C.COMPARISON_TABLES:
        urls.append((SITE_URL + url_for('compare_one', slug=c['slug']), 'monthly',
                     '0.7', SEARCH_META_CHANGED))

    # One URL per document format, for the same reason: "board resolution
    # format" is a real query and /templates alone could never answer it.
    for r in format_rows:
        urls.append((SITE_URL + url_for('format_page', slug=r['slug']), 'yearly', '0.6',
                     freshest(r['updated_at'] or '', SEARCH_META_CHANGED)))

    # Resolution library pages
    for rtype in DOC_LIST_TITLE:
        urls.append((SITE_URL + url_for('resolutions', rtype=rtype), 'monthly', '0.6', None))

    # Landmark judgment briefs
    for jd in C.JUDGMENTS:
        urls.append((SITE_URL + url_for('judgment', slug=jd['slug']), 'monthly', '0.6',
                     freshest(JUDGMENTS_PUBLISHED, SEARCH_META_CHANGED)))

    # Published articles. The body's own date, unless the title and description
    # Google displays were rewritten more recently — the page a crawler fetches
    # changed on that later date even though the article text did not.
    for r in article_rows:
        urls.append((SITE_URL + url_for('article', slug=r['slug']), 'monthly', '0.8',
                     freshest(r['updated_at'] or '', SEARCH_META_CHANGED)))

    xml = ['<?xml version="1.0" encoding="UTF-8"?>',
           '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">']
    for loc, freq, prio, lastmod in urls:
        xml.append('  <url>')
        xml.append(f'    <loc>{loc}</loc>')
        if lastmod:
            xml.append(f'    <lastmod>{lastmod}</lastmod>')
        xml.append(f'    <changefreq>{freq}</changefreq>')
        xml.append(f'    <priority>{prio}</priority>')
        xml.append('  </url>')
    xml.append('</urlset>')
    return app.response_class('\n'.join(xml), mimetype='application/xml')


@app.route('/llms.txt')
def llms_txt():
    """AI-search discovery file (llmstxt.org): a clean, linked index of the
    site's guides so ChatGPT/Perplexity/AI Overviews can find and cite them."""
    from itertools import groupby
    db = get_db()
    rows = db.execute(
        'SELECT slug,title,summary,category FROM articles WHERE published=1 '
        'ORDER BY category,title'
    ).fetchall()
    db.close()
    out = [
        '# Law Minded',
        '',
        "> India's plain-English legal and compliance platform — corporate "
        'compliance, labour law, consumer rights and constitutional law, '
        'explained simply for founders, professionals and citizens.',
        '',
        '## Key pages',
        f'- [Knowledge Hub]({SITE_URL}/blogs): every guide, by topic',
        f'- [Free templates]({SITE_URL}/templates): ready-to-use legal document formats',
        f'- [Landmark judgments]({SITE_URL}/judgments): plain-English case summaries',
        f'- [FAQ]({SITE_URL}/faq)',
        '',
    ]
    for cat, items in groupby(rows, key=lambda r: r['category']):
        out.append(f'## {CATEGORY_MAP.get(cat, (cat or "Guides").title())}')
        for r in items:
            summ = ' '.join((r['summary'] or '').split())[:140]
            out.append(f'- [{r["title"]}]({SITE_URL}/article/{r["slug"]})'
                       + (f': {summ}' if summ else ''))
        out.append('')
    return app.response_class('\n'.join(out), mimetype='text/plain; charset=utf-8')


@app.errorhandler(404)
def not_found(e):
    return render_template('404.html'), 404


# ─── Admin Routes ────────────────────────────────────────────────────────────

@app.route('/admin/login', methods=['GET', 'POST'])
@limiter.limit('10 per minute; 40 per hour', methods=['POST'])
def admin_login():
    if request.method == 'POST':
        username = request.form.get('username', '')
        password = request.form.get('password', '')
        if check_admin(username, password):
            session.clear()                      # rotate session on login (anti-fixation)
            session.permanent = True             # subject to PERMANENT_SESSION_LIFETIME
            session['admin_logged_in'] = True
            return redirect(url_for('admin_dashboard'))
        flash('Incorrect username or password.', 'error')
    return render_template('admin/login.html')


@app.route('/admin/password', methods=['GET', 'POST'])
@admin_required
def admin_password():
    if request.method == 'POST':
        current = request.form.get('current', '')
        new_username = request.form.get('username', '').strip()
        new = request.form.get('new', '')
        confirm = request.form.get('confirm', '')
        cur_hash = _admin_hash()
        if not cur_hash or not check_password_hash(cur_hash, current):
            flash('Your current password is incorrect.', 'error')
        elif new_username and not re.fullmatch(r'[A-Za-z0-9._-]{3,40}', new_username):
            flash('Username must be 3–40 characters (letters, numbers, . _ - only).', 'error')
        elif len(new) < 8:
            flash('New password must be at least 8 characters.', 'error')
        elif new != confirm:
            flash('New password and confirmation do not match.', 'error')
        else:
            set_admin_credentials(username=new_username or get_admin_username(), password=new)
            flash('Admin credentials updated successfully.', 'success')
            return redirect(url_for('admin_dashboard'))
    return render_template('admin/password.html', admin_username=get_admin_username())


@app.route('/admin/logout')
def admin_logout():
    session.clear()
    return redirect(url_for('admin_login'))


@app.route('/admin')
@admin_required
def admin_dashboard():
    db = get_db()
    stats = {
        'articles': db.execute('SELECT COUNT(*) FROM articles').fetchone()[0],
        'published': db.execute('SELECT COUNT(*) FROM articles WHERE published=1').fetchone()[0],
        'subscribers': db.execute('SELECT COUNT(*) FROM subscribers').fetchone()[0],
        'messages': db.execute('SELECT COUNT(*) FROM contact_messages').fetchone()[0],
    }
    recent_messages = db.execute(
        'SELECT * FROM contact_messages ORDER BY created_at DESC LIMIT 5'
    ).fetchall()
    db.close()
    return render_template('admin/dashboard.html', stats=stats, recent_messages=recent_messages)


@app.route('/admin/automation')
@admin_required
def admin_automation():
    """What the writing automation is doing. Two sources: this database, which
    knows what is drafted and scheduled, and a JSON snapshot the writer box
    pushes here every quarter hour. The snapshot is deliberately not fetched live
    — a writer box that is slow or down should not hang the admin panel of a site
    that does not otherwise need it."""
    status, status_error = {}, None
    snap = os.path.join(os.path.dirname(DB_PATH), 'writer-status.json')
    try:
        with open(snap) as f:
            status = json.load(f)
    except FileNotFoundError:
        status_error = 'No report from the writer box yet.'
    except (OSError, ValueError) as e:
        status_error = f'Could not read the writer box report: {e}'

    if status.get('generated_at'):
        try:
            age = datetime.now(IST) - datetime.fromisoformat(status['generated_at'])
            status['age_minutes'] = int(age.total_seconds() // 60)
        except ValueError:
            status['age_minutes'] = None

    db = get_db()
    drafts = db.execute(
        'SELECT slug, title, publish_on, created_at FROM articles '
        'WHERE published=0 ORDER BY publish_on IS NULL, publish_on, created_at'
    ).fetchall()
    recent = db.execute(
        'SELECT slug, title, created_at FROM articles WHERE published=1 '
        'ORDER BY created_at DESC LIMIT 8').fetchall()
    published = db.execute('SELECT COUNT(*) FROM articles WHERE published=1').fetchone()[0]
    db.close()

    return render_template('admin/automation.html', status=status,
                           status_error=status_error, drafts=drafts,
                           recent=recent, published=published,
                           draft_url=draft_url, today=ist_today())


@app.route('/admin/articles')
@admin_required
def admin_articles():
    db = get_db()
    articles = db.execute('SELECT * FROM articles ORDER BY created_at DESC').fetchall()
    db.close()
    return render_template('admin/articles.html', articles=articles)


@app.route('/admin/articles/new', methods=['GET', 'POST'])
@admin_required
def admin_article_new():
    if request.method == 'POST':
        title = request.form.get('title', '').strip()
        category = request.form.get('category', '').strip()
        act = request.form.get('act', '').strip()
        read_time = request.form.get('read_time', '').strip()
        summary = request.form.get('summary', '').strip()
        seo_title = request.form.get('seo_title', '').strip()
        content_html = sanitize_html(request.form.get('content', '').strip())
        published = 1 if request.form.get('published') else 0
        slug = slugify(title)

        db = get_db()
        base_slug = slug
        i = 1
        while db.execute('SELECT id FROM articles WHERE slug=?', (slug,)).fetchone():
            slug = f'{base_slug}-{i}'
            i += 1
        db.execute(
            'INSERT INTO articles (title, slug, category, act, read_time, summary, seo_title, content, published) VALUES (?,?,?,?,?,?,?,?,?)',
            (title, slug, category, act, read_time, summary, seo_title, content_html, published)
        )
        db.commit()
        db.close()
        flash('Article created successfully.', 'success')
        return redirect(url_for('admin_articles'))

    return render_template('admin/article_edit.html', article=None)


@app.route('/admin/articles/<int:article_id>/edit', methods=['GET', 'POST'])
@admin_required
def admin_article_edit(article_id):
    db = get_db()
    art = db.execute('SELECT * FROM articles WHERE id=?', (article_id,)).fetchone()
    if not art:
        db.close()
        abort(404)

    if request.method == 'POST':
        title = request.form.get('title', '').strip()
        category = request.form.get('category', '').strip()
        act = request.form.get('act', '').strip()
        read_time = request.form.get('read_time', '').strip()
        summary = request.form.get('summary', '').strip()
        seo_title = request.form.get('seo_title', '').strip()
        content_html = sanitize_html(request.form.get('content', '').strip())
        published = 1 if request.form.get('published') else 0
        db.execute(
            '''UPDATE articles SET title=?, category=?, act=?, read_time=?, summary=?,
               seo_title=?, content=?, published=?, updated_at=CURRENT_TIMESTAMP WHERE id=?''',
            (title, category, act, read_time, summary, seo_title, content_html, published, article_id)
        )
        db.commit()
        db.close()
        flash('Article updated successfully.', 'success')
        return redirect(url_for('admin_articles'))

    db.close()
    return render_template('admin/article_edit.html', article=art)


@app.route('/admin/articles/<int:article_id>/delete', methods=['POST'])
@admin_required
def admin_article_delete(article_id):
    db = get_db()
    db.execute('DELETE FROM articles WHERE id=?', (article_id,))
    db.commit()
    db.close()
    flash('Article deleted.', 'success')
    return redirect(url_for('admin_articles'))


@app.route('/admin/subscribers')
@admin_required
def admin_subscribers():
    db = get_db()
    subscribers = db.execute('SELECT * FROM subscribers ORDER BY created_at DESC').fetchall()
    db.close()
    return render_template('admin/subscribers.html', subscribers=subscribers)


@app.route('/admin/subscribers/add', methods=['POST'])
@admin_required
def admin_subscriber_add():
    name = request.form.get('name', '').strip()
    email = request.form.get('email', '').strip()
    if not email or not VALID_EMAIL_RE.match(email):
        flash('Please enter a valid email address.', 'error')
        return redirect(url_for('admin_subscribers'))
    db = get_db()
    try:
        db.execute('INSERT INTO subscribers (name, email) VALUES (?,?)', (name, email))
        db.commit()
        flash(f'Subscriber {email} added.', 'success')
    except sqlite3.IntegrityError:
        flash(f'{email} is already subscribed.', 'error')
    db.close()
    return redirect(url_for('admin_subscribers'))


@app.route('/admin/subscribers/<int:sub_id>/delete', methods=['POST'])
@admin_required
def admin_subscriber_delete(sub_id):
    db = get_db()
    db.execute('DELETE FROM subscribers WHERE id=?', (sub_id,))
    db.commit()
    db.close()
    flash('Subscriber removed.', 'success')
    return redirect(url_for('admin_subscribers'))


@app.route('/admin/subscribers/export')
@admin_required
def admin_subscribers_export():
    """Download all subscribers as a real .xlsx (opens directly in Excel)."""
    db = get_db()
    rows = db.execute(
        'SELECT name, email, created_at FROM subscribers ORDER BY created_at DESC'
    ).fetchall()
    db.close()

    from openpyxl import Workbook
    from openpyxl.styles import Font
    wb = Workbook()
    ws = wb.active
    ws.title = 'Subscribers'
    ws.append(['Name', 'Email', 'Subscribed On'])
    for col in ('A1', 'B1', 'C1'):
        ws[col].font = Font(bold=True)
    for r in rows:
        ws.append([r['name'] or '', r['email'], (str(r['created_at'])[:19] if r['created_at'] else '')])
    ws.column_dimensions['A'].width = 24
    ws.column_dimensions['B'].width = 36
    ws.column_dimensions['C'].width = 22
    ws.freeze_panes = 'A2'

    bio = BytesIO()
    wb.save(bio)
    bio.seek(0)
    fname = f'lawminded-subscribers-{date.today().isoformat()}.xlsx'
    return send_file(
        bio, as_attachment=True, download_name=fname,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )


@app.route('/admin/messages')
@admin_required
def admin_messages():
    db = get_db()
    messages = db.execute('SELECT * FROM contact_messages ORDER BY created_at DESC').fetchall()
    db.close()
    return render_template('admin/messages.html', messages=messages)


# ─── Admin: Documents (templates & resolutions) ──────────────────────────────

RESOLUTION_TYPES = ('board',)


def _doc_home(doc_type):
    """Which admin listing a document belongs to, based on its type."""
    return 'admin_resolutions' if doc_type in RESOLUTION_TYPES else 'admin_documents'


@app.route('/admin/documents')
@admin_required
def admin_documents():
    db = get_db()
    docs = db.execute(
        "SELECT * FROM documents WHERE doc_type='template' ORDER BY sort_order, id"
    ).fetchall()
    db.close()
    return render_template(
        'admin/documents.html',
        grouped={'template': list(docs)},
        doc_types={'template': DOC_TYPES['template']},
        admin_section='templates',
        page_title='Templates',
        page_intro='Manage your downloadable document templates. Changes appear on the site instantly.',
        new_label='+ New Template',
        new_default_type='template',
    )


@app.route('/admin/resolutions')
@admin_required
def admin_resolutions():
    db = get_db()
    placeholders = ','.join('?' * len(RESOLUTION_TYPES))
    docs = db.execute(
        f"SELECT * FROM documents WHERE doc_type IN ({placeholders}) "
        "ORDER BY doc_type, sort_order, id", RESOLUTION_TYPES
    ).fetchall()
    db.close()
    grouped = {k: [] for k in RESOLUTION_TYPES}
    for d in docs:
        grouped.setdefault(d['doc_type'], []).append(d)
    return render_template(
        'admin/documents.html',
        grouped=grouped,
        doc_types={k: DOC_TYPES[k] for k in RESOLUTION_TYPES},
        admin_section='resolutions',
        page_title='Resolutions',
        page_intro='Manage your Board resolutions. Changes appear on the site instantly.',
        new_label='+ New Resolution',
        new_default_type='board',
    )


def _save_document(form, doc_id=None):
    doc_type = form.get('doc_type', '').strip()
    title = form.get('title', '').strip()
    icon = form.get('icon', '').strip() or '📄'
    description = form.get('description', '').strip()
    tags = ','.join([t.strip() for t in form.get('tags', '').split(',') if t.strip()])
    body = form.get('body', '').strip()
    if doc_type not in DOC_TYPES or not title or not body:
        return None, 'Type, title and body are required.'
    db = get_db()
    base = slugify(title) or 'document'
    slug = base
    i = 1
    while True:
        clash = db.execute(
            'SELECT id FROM documents WHERE doc_type=? AND slug=? AND id IS NOT ?',
            (doc_type, slug, doc_id)
        ).fetchone()
        if not clash:
            break
        slug = f'{base}-{i}'; i += 1
    if doc_id:
        db.execute(
            'UPDATE documents SET doc_type=?, slug=?, icon=?, title=?, description=?, tags=?, body=?, '
            'updated_at=CURRENT_TIMESTAMP WHERE id=?',
            (doc_type, slug, icon, title, description, tags, body, doc_id)
        )
    else:
        order = db.execute('SELECT COALESCE(MAX(sort_order),0)+1 FROM documents WHERE doc_type=?',
                           (doc_type,)).fetchone()[0]
        db.execute(
            'INSERT INTO documents (doc_type, slug, icon, title, description, tags, body, sort_order) '
            'VALUES (?,?,?,?,?,?,?,?)',
            (doc_type, slug, icon, title, description, tags, body, order)
        )
    db.commit()
    db.close()
    return True, None


@app.route('/admin/documents/new', methods=['GET', 'POST'])
@admin_required
def admin_document_new():
    if request.method == 'POST':
        ok, err = _save_document(request.form)
        if ok:
            flash('Document created.', 'success')
            return redirect(url_for(_doc_home(request.form.get('doc_type', ''))))
        flash(err, 'error')
    default_type = request.form.get('doc_type') or request.args.get('type') or 'template'
    return render_template('admin/document_edit.html', doc=None, doc_types=DOC_TYPES,
                           default_type=default_type,
                           admin_section=('resolutions' if default_type in RESOLUTION_TYPES else 'templates'))


@app.route('/admin/documents/<int:doc_id>/edit', methods=['GET', 'POST'])
@admin_required
def admin_document_edit(doc_id):
    db = get_db()
    doc = db.execute('SELECT * FROM documents WHERE id=?', (doc_id,)).fetchone()
    db.close()
    if not doc:
        abort(404)
    if request.method == 'POST':
        ok, err = _save_document(request.form, doc_id=doc_id)
        if ok:
            flash('Document updated.', 'success')
            return redirect(url_for(_doc_home(request.form.get('doc_type', doc['doc_type']))))
        flash(err, 'error')
    return render_template('admin/document_edit.html', doc=doc, doc_types=DOC_TYPES,
                           default_type=doc['doc_type'],
                           admin_section=('resolutions' if doc['doc_type'] in RESOLUTION_TYPES else 'templates'))


@app.route('/admin/documents/<int:doc_id>/delete', methods=['POST'])
@admin_required
def admin_document_delete(doc_id):
    db = get_db()
    row = db.execute('SELECT doc_type FROM documents WHERE id=?', (doc_id,)).fetchone()
    db.execute('DELETE FROM documents WHERE id=?', (doc_id,))
    db.commit()
    db.close()
    flash('Document deleted.', 'success')
    return redirect(url_for(_doc_home(row['doc_type']) if row else 'admin_documents'))


# ─── Admin: Document Formats (uploaded Word files) ───────────────────────────
ALLOWED_FORMAT_EXTS = {'.docx'}


@app.route('/admin/formats')
@admin_required
def admin_formats():
    db = get_db()
    rows = db.execute('SELECT * FROM formats ORDER BY sort_order, id').fetchall()
    db.close()
    by_cat = {}
    for r in rows:
        by_cat.setdefault(r['category'], []).append(r)
    grouped = {}
    for name in F.CATEGORY_NAMES:
        if name in by_cat:
            grouped[name] = by_cat[name]
    for name, items in by_cat.items():
        grouped.setdefault(name, items)
    return render_template('admin/formats.html', grouped=grouped, categories=F.CATEGORY_NAMES)


def _save_format(form, files, fmt_id=None, existing=None):
    category = form.get('category', '').strip()
    title = form.get('title', '').strip()
    description = form.get('description', '').strip()
    if not category or not title:
        return None, 'Category and title are required.'
    filename = existing['filename'] if existing else None
    upload = files.get('docfile')
    if upload and upload.filename:
        if os.path.splitext(upload.filename)[1].lower() not in ALLOWED_FORMAT_EXTS:
            return None, 'Only Word (.docx) files can be uploaded.'
        base = secure_filename(os.path.splitext(upload.filename)[0]) or 'document'
        fname = f'{base}.docx'
        os.makedirs(FORMATS_DIR, exist_ok=True)
        i = 1
        while os.path.exists(os.path.join(FORMATS_DIR, fname)):
            fname = f'{base}-{i}.docx'
            i += 1
        upload.save(os.path.join(FORMATS_DIR, fname))
        filename = fname
    if not filename:
        return None, 'Please choose a Word (.docx) file to upload.'
    db = get_db()
    base_slug = slugify(title) or 'document'
    slug = base_slug
    i = 1
    while db.execute('SELECT id FROM formats WHERE slug=? AND id IS NOT ?', (slug, fmt_id)).fetchone():
        slug = f'{base_slug}-{i}'
        i += 1
    if fmt_id:
        db.execute(
            'UPDATE formats SET category=?, slug=?, title=?, description=?, filename=?, '
            'updated_at=CURRENT_TIMESTAMP WHERE id=?',
            (category, slug, title, description, filename, fmt_id)
        )
    else:
        order = db.execute('SELECT COALESCE(MAX(sort_order),0)+1 FROM formats').fetchone()[0]
        db.execute(
            'INSERT INTO formats (category, slug, title, description, filename, sort_order) '
            'VALUES (?,?,?,?,?,?)',
            (category, slug, title, description, filename, order)
        )
    db.commit()
    db.close()
    _render_format_preview.cache_clear()
    return True, None


@app.route('/admin/formats/new', methods=['GET', 'POST'])
@admin_required
def admin_format_new():
    if request.method == 'POST':
        ok, err = _save_format(request.form, request.files)
        if ok:
            flash('Document uploaded.', 'success')
            return redirect(url_for('admin_formats'))
        flash(err, 'error')
    return render_template('admin/format_edit.html', fmt=None, categories=F.CATEGORY_NAMES)


@app.route('/admin/formats/<int:fmt_id>/edit', methods=['GET', 'POST'])
@admin_required
def admin_format_edit(fmt_id):
    db = get_db()
    fmt = db.execute('SELECT * FROM formats WHERE id=?', (fmt_id,)).fetchone()
    db.close()
    if not fmt:
        abort(404)
    if request.method == 'POST':
        ok, err = _save_format(request.form, request.files, fmt_id=fmt_id, existing=fmt)
        if ok:
            flash('Document updated.', 'success')
            return redirect(url_for('admin_formats'))
        flash(err, 'error')
    return render_template('admin/format_edit.html', fmt=fmt, categories=F.CATEGORY_NAMES)


@app.route('/admin/formats/<int:fmt_id>/delete', methods=['POST'])
@admin_required
def admin_format_delete(fmt_id):
    db = get_db()
    db.execute('DELETE FROM formats WHERE id=?', (fmt_id,))
    db.commit()
    db.close()
    _render_format_preview.cache_clear()
    flash('Document deleted.', 'success')
    return redirect(url_for('admin_formats'))


# ─── App Init ────────────────────────────────────────────────────────────────

with app.app_context():
    init_db()
    seed_articles()
    seed_documents()
    seed_formats()
    # After the seeders, never before: the content fixes edit seeded rows, and
    # running them against an empty table stamps them as applied for good.
    apply_content_migrations()

if __name__ == '__main__':
    # PORT lets the dev preview pick a free port; defaults to 8000 locally.
    app.run(host='0.0.0.0', port=int(os.getenv('PORT', '8000')), debug=not IS_PROD)
